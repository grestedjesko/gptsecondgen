import io
import os
import textwrap
import tempfile
from typing import Optional

import pandas as pd
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from typing import Union, IO

try:
    import docx  # python-docx
except ImportError:
    docx = None


class SimpleFileToPDF:
    """
    Лёгкий конвертер без LibreOffice:
      - .docx: текст параграфов -> PDF
      - .xlsx/.xls/.csv: таблицы -> PDF
      - .txt: текст -> PDF
    Не поддерживает .doc (старая бинарная версия Word) без внешних утилит.
    """

    SUPPORTED = {".docx", ".xlsx", ".xls", ".csv", ".txt"}

    def __init__(self, page_size=A4, table_landscape=True, margin_mm=15):
        self.page_size = page_size
        self.table_landscape = table_landscape
        self.margin = margin_mm * mm
        self.styles = getSampleStyleSheet()

    def _to_input_bytes(self, data: Union[bytes, bytearray, memoryview, IO[bytes]]) -> bytes:
        """Приводит вход к bytes: поддерживает bytes/bytearray/memoryview/файлоподобные (BytesIO)."""
        if isinstance(data, bytes):
            return data
        if isinstance(data, (bytearray, memoryview)):
            return bytes(data)
        # файлоподобный объект
        read = getattr(data, "read", None)
        if callable(read):
            # Некоторые потоки нужно перематывать
            try:
                pos = getattr(data, "tell", lambda: None)()
                getattr(data, "seek", lambda *_: None)(0)
            except Exception:
                pass
            content = read()
            # вернёмся на место, если получилось
            try:
                if pos is not None:
                    getattr(data, "seek", lambda *_: None)(pos)
            except Exception:
                pass
            if not isinstance(content, (bytes, bytearray, memoryview)):
                raise TypeError(f"read() вернул {type(content)!r}, ожидались bytes-подобные")
            return bytes(content)
        raise TypeError(f"Ожидались bytes/bytearray/memoryview или файлоподобный объект, получено {type(data)!r}")

    # ---------- Публичный API ----------

    def convert(self, file_bytes, filename: Optional[str] = None, mime_type: Optional[str] = None) -> bytes:
        ext = self._detect_ext(filename, mime_type)
        if ext == '.pdf':
            return self._to_input_bytes(file_bytes)

        if ext not in self.SUPPORTED:
            raise ValueError(f"Неподдерживаемое расширение: {ext}. Поддерживаются: {', '.join(sorted(self.SUPPORTED))}")

        file_bytes = self._to_input_bytes(file_bytes)

        with tempfile.TemporaryDirectory() as td:
            src_path = os.path.join(td, f"input{ext}")
            with open(src_path, "wb") as f:
                f.write(file_bytes)

            if ext == ".docx":
                pdf = self._docx_to_pdf(src_path)
            elif ext in (".xlsx", ".xls"):
                pdf = self._excel_to_pdf(src_path)
            elif ext == ".csv":
                pdf = self._csv_to_pdf(src_path)
            elif ext == ".txt":
                pdf = self._txt_to_pdf(src_path)
            else:
                raise ValueError("Неподдерживаемый тип")

        return pdf

    # ---------- Конверторы ----------

    def _docx_to_pdf(self, path: str) -> bytes:
        if docx is None:
            raise RuntimeError("Не установлен python-docx: pip install python-docx")

        doc = docx.Document(path)
        buf = io.BytesIO()
        doc_tpl = SimpleDocTemplate(
            buf, pagesize=self.page_size,
            leftMargin=self.margin, rightMargin=self.margin,
            topMargin=self.margin, bottomMargin=self.margin
        )
        elements = []
        hstyle = self.styles["Heading2"]
        pstyle = self.styles["BodyText"]

        # Пробуем простую семантику: если стиль параграфа содержит "Heading", считаем заголовком
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                elements.append(Spacer(1, 6))
                continue
            style_name = (para.style.name or "").lower()
            style = hstyle if "heading" in style_name else pstyle
            elements.append(Paragraph(self._escape(text), style))
            elements.append(Spacer(1, 6))

        # Таблицы тоже перекинем как таблицы
        for t in doc.tables:
            data = [[self._escape(cell.text.strip()) for cell in row.cells] for row in t.rows]
            if not data:
                continue
            tbl = self._make_table(data)
            elements.append(tbl)
            elements.append(Spacer(1, 10))

        doc_tpl.build(elements)
        return buf.getvalue()

    def _excel_to_pdf(self, path: str) -> bytes:
        # читаем все листы
        xls = pd.ExcelFile(path)
        buf = io.BytesIO()

        # Для широких таблиц перевернём страницу в альбомную
        pagesize = landscape(self.page_size) if self.table_landscape else self.page_size

        c = canvas.Canvas(buf, pagesize=pagesize)
        width, height = pagesize

        for i, sheet in enumerate(xls.sheet_names):
            df = xls.parse(sheet_name=sheet)
            self._draw_dataframe_as_tables(c, df, page_title=f"{os.path.basename(path)} — {sheet}", width=width, height=height)
            if i < len(xls.sheet_names) - 1:
                c.showPage()
        c.save()
        return buf.getvalue()

    def _csv_to_pdf(self, path: str) -> bytes:
        df = pd.read_csv(path)
        return self._df_to_pdf(df, title=os.path.basename(path))

    def _txt_to_pdf(self, path: str) -> bytes:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=self.page_size)
        width, height = self.page_size

        x = self.margin
        y = height - self.margin
        line_h = 12
        max_width_chars = 100  # грубая оценка для wrap

        for paragraph in text.splitlines():
            lines = textwrap.wrap(paragraph, width=max_width_chars) or [""]
            for ln in lines:
                if y < self.margin + line_h:
                    c.showPage()
                    y = height - self.margin
                c.drawString(x, y, ln)
                y -= line_h

        c.save()
        return buf.getvalue()

    # ---------- Вспомогательное ----------

    def _df_to_pdf(self, df: pd.DataFrame, title: Optional[str] = None) -> bytes:
        buf = io.BytesIO()
        pagesize = landscape(self.page_size) if self.table_landscape else self.page_size
        c = canvas.Canvas(buf, pagesize=pagesize)
        width, height = pagesize
        self._draw_dataframe_as_tables(c, df, page_title=title, width=width, height=height)
        c.save()
        return buf.getvalue()

    def _draw_dataframe_as_tables(self, c: canvas.Canvas, df: pd.DataFrame, page_title: Optional[str], width, height):
        # поля
        x0 = self.margin
        y = height - self.margin

        # Заголовок
        if page_title:
            c.setFont("Helvetica-Bold", 12)
            c.drawString(x0, y, page_title)
            y -= 16

        # конвертируем DataFrame в список списков, с заголовками
        header = [str(col) for col in df.columns]
        data_rows = df.astype(str).values.tolist()
        data = [header] + data_rows

        # Порционируем по ~40 строк на таблицу (зависит от шрифта/полей)
        rows_per_page = 40
        start = 0
        while start < len(data):
            chunk = data[start:start + rows_per_page]
            table = self._make_table(chunk, max_width=width - 2 * self.margin)
            w, h = table.wrapOn(c, width, height)
            if y - h < self.margin:
                c.showPage()
                y = height - self.margin
                if page_title:
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(x0, y, page_title)
                    y -= 16
            table.drawOn(c, x0, y - h)
            y -= (h + 10)
            start += rows_per_page

    def _make_table(self, data, max_width=None):
        # ширины колонок по содержимому (грубая оценка)
        col_count = max(len(r) for r in data)
        col_width = (max_width or (self.page_size[0] - 2 * self.margin)) / max(col_count, 1)
        table = Table(data, colWidths=[col_width] * col_count, repeatRows=1)
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return table

    def _escape(self, text: str) -> str:
        # Простейшая экранизация для Paragraph (минимум XML-символов)
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

    def _detect_ext(self, filename: Optional[str], mime: Optional[str]) -> str:
        if filename and "." in filename:
            return os.path.splitext(filename)[1].lower()
        # Если имя не дали — попробуем по mime
        import mimetypes
        ext = mimetypes.guess_extension(mime or "") or ""
        return ext.lower()
